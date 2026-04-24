import os
import re
import csv
import subprocess
from dataclasses import dataclass, asdict
from typing import Dict, List, Optional, Tuple

try:
    from docx import Document  # type: ignore
except ImportError:
    Document = None


SOURCE_FOLDER = "/Users/local/Desktop/POLISY_DOCX"
OUTPUT_CSV = "/Users/local/Desktop/oc_policies_extracted_v2.csv"
MAX_FILES: Optional[int] = None


@dataclass
class PolicyRecord:
    source_file: str = ""
    product_type: str = ""
    insurer: str = ""
    policyholder: str = ""
    insured: str = ""
    broker: str = ""
    risk_code: str = ""
    insured_activity: str = ""
    scope_of_insurance: str = ""
    territorial_scope: str = ""
    insurance_period: str = ""
    sum_guaranteed: str = ""
    sum_guaranteed_notes: str = ""
    deductible_main: str = ""
    deductible_notes: str = ""
    turnover: str = ""
    premium_rate: str = ""
    premium: str = ""
    law_and_jurisdiction: str = ""
    has_office_liability: str = ""
    has_documents_loss: str = ""
    has_subcontractors_cover: str = ""
    has_subcontractors_regress: str = ""
    covers_operations: str = ""
    covers_product_liability: str = ""
    covers_completed_operations: str = ""
    covers_professional_liability: str = ""
    covers_pure_financial_loss: str = ""
    has_employer_liability: str = ""
    has_property_under_control: str = ""
    has_environment_damage: str = ""
    has_travel_clause: str = ""
    has_vehicles_clause: str = ""
    has_leased_property: str = ""
    has_extended_product_clauses: str = ""
    has_construction_clause: str = ""
    has_perimeter_clause: str = ""
    has_vibration_clause: str = ""
    has_construction_machinery_clause: str = ""
    has_design_prof_liability: str = ""
    has_building_damage_cover: str = ""
    has_gradual_damage_cover: str = ""
    has_cost_overrun_extension: str = ""
    has_deadline_overrun_extension: str = ""
    has_contractual_penalties_regress: str = ""
    profession_subtype: str = ""
    has_excess_layer: str = ""
    attachment_point: str = ""
    documents_limit_rule: str = ""
    classification_reasons: str = ""
    data_quality_flag: str = ""


def normalize(text: str) -> str:
    text = text.replace("\xa0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\r", "\n", text)
    text = re.sub(r"\n{2,}", "\n\n", text)
    return text.strip()


def compact_whitespace(s: str) -> str:
    return re.sub(r"\s+", " ", s).strip()


def lower_pl(text: str) -> str:
    return normalize(text).lower()


def yesno(flag: bool) -> str:
    return "yes" if flag else "no"


def read_docx(path: str) -> str:
    if Document is None:
        return ""
    try:
        doc = Document(path)
        parts: List[str] = []
        for p in doc.paragraphs:
            if p.text:
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                row_parts = []
                for cell in row.cells:
                    txt = cell.text.strip()
                    if txt:
                        row_parts.append(txt)
                if row_parts:
                    parts.append(" ".join(row_parts))
        return "\n".join(parts)
    except Exception:
        return ""


def read_doc(path: str) -> str:
    try:
        result = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", path],
            capture_output=True,
            text=True,
            check=False,
        )
        return result.stdout or ""
    except Exception:
        return ""


def read_file_text(path: str) -> str:
    lower = path.lower()
    if lower.endswith(".docx"):
        return read_docx(path)
    if lower.endswith(".doc"):
        return read_doc(path)
    return ""


def find_first(patterns: List[str], text: str, flags: int = re.IGNORECASE | re.MULTILINE) -> str:
    for pattern in patterns:
        m = re.search(pattern, text, flags)
        if m:
            if m.lastindex:
                return compact_whitespace(m.group(1))
            return compact_whitespace(m.group(0))
    return ""


def find_all_keywords(text_lower: str, keywords: List[str]) -> bool:
    return any(k in text_lower for k in keywords)


def extract_section(text: str, start_keyword: str, end_keywords: List[str]) -> str:
    text_norm = normalize(text)
    text_lower = text_norm.lower()
    start = text_lower.find(start_keyword.lower())
    if start == -1:
        return ""
    content_start = start + len(start_keyword)
    end = len(text_norm)
    for kw in end_keywords:
        pos = text_lower.find(kw.lower(), content_start)
        if pos != -1 and pos < end:
            end = pos
    return compact_whitespace(text_norm[content_start:end])


def classify_policy(text: str, filename: str) -> Tuple[str, List[str]]:
    tl = lower_pl(text)
    reasons: List[str] = []
    score = {
        "oc_ogolne": 0,
        "oc_budowlane": 0,
        "oc_architekci": 0,
        "oc_zawodowe": 0,
    }

    zawodowe_terms = [
        "ubezpieczenia odpowiedzialności cywilnej zawodowej",
        "czynności zawodowe brokera ubezpieczeniowego",
        "świadczenie pomocy prawnej",
        "doradztwo podatkowe",
        "konsulting",
        "agencji reklamowych",
        "zarządzaniu nieruchomościami",
        "zarządców nieruchomości",
    ]
    if find_all_keywords(tl, zawodowe_terms):
        score["oc_zawodowe"] += 8
        reasons.append("słowa kluczowe OC zawodowe")

    arch_terms = [
        "oc architektów",
        "zawodu projektanta",
        "osób wykonujących zawody techniczne",
        "sprawowanie nadzoru autorskiego",
        "projektowanie z zakresu",
    ]
    if find_all_keywords(tl, arch_terms):
        score["oc_architekci"] += 8
        reasons.append("słowa kluczowe OC architektów/projektantów")

    bud_terms = [
        "przedsiębiorstw budowlanych",
        "prac budowlanych",
        "robót budowlanych",
        "montażowych",
        "klauzula promienia",
        "wibracji, osłabienia elementów nośnych",
        "samobieżne maszyny budowlane",
    ]
    if find_all_keywords(tl, bud_terms):
        score["oc_budowlane"] += 8
        reasons.append("słowa kluczowe OC budowlanego")

    ogolne_terms = [
        "odpowiedzialność cywilna z tytułu prowadzonej działalności",
        "posiadanego mienia",
        "odpowiedzialność za produkt",
        "wykonane usługi",
        "odpowiedzialność cywilna ogólna",
    ]
    if find_all_keywords(tl, ogolne_terms):
        score["oc_ogolne"] += 8
        reasons.append("słowa kluczowe OC ogólnego")

    if "zawodowa odpowiedzialność cywilna" in tl:
        score["oc_zawodowe"] += 3
    if "wykonywania zawodu projektanta" in tl:
        score["oc_architekci"] += 4
        score["oc_zawodowe"] -= 1
    if "odpowiedzialność za produkt" in tl:
        score["oc_ogolne"] += 3
        score["oc_budowlane"] += 1
    if "czynności zawodowe" in tl:
        score["oc_zawodowe"] += 2
    if "sprawowanie nadzoru autorskiego" in tl:
        score["oc_architekci"] += 3
    if "broker ubezpieczeniowy" in tl:
        score["oc_zawodowe"] += 4
    if "świadczenie pomocy prawnej" in tl:
        score["oc_zawodowe"] += 4
    if "projektowanie" in tl and "budowlanych" in tl:
        score["oc_architekci"] += 2

    winner = max(score.items(), key=lambda kv: kv[1])[0]
    reasons.append("scores=" + str(score))
    return winner, reasons


def extract_common_fields(text: str, filename: str, product_type: str, reasons: List[str]) -> PolicyRecord:
    t = normalize(text)
    tl = lower_pl(text)
    record = PolicyRecord(
        source_file=filename,
        product_type=product_type,
        classification_reasons=" | ".join(reasons),
    )

    record.insurer = find_first([
        r"Ubezpieczyciel(?:\s*/\s*Insurer)?\s+(.*?)\s+(?:Ubezpieczający|Policyholder|Ubezpieczający i Ubezpieczony|ubezpieczający)",
    ], t)

    record.policyholder = find_first([
        r"Ubezpieczający i Ubezpieczony\s+(.*?)\s+Ubezpieczona działalność",
        r"Ubezpieczający(?:\s*/\s*Policyholder)?\s+(.*?)\s+(?:Ubezpieczony|Ubezpieczeni|Insured|Broker|Agent)",
        r"ubezpieczający\s+(.*?)\s+ubezpieczony",
    ], t)

    record.insured = find_first([
        r"Ubezpieczony(?:\s*/\s*Insured)?\s+(.*?)\s+(?:Broker|Agent|Ubezpieczona działalność|Zakres ubezpieczenia)",
        r"Ubezpieczeni(?:\s*/\s*Insureds)?\s+(.*?)\s+(?:Broker|Ubezpieczona działalność|Zakres ubezpieczenia)",
        r"ubezpieczony\s+(.*?)\s+broker",
    ], t)

    record.broker = find_first([
        r"Broker\s+(.*?)\s+(?:Ubezpieczona działalność|Kod ryzyka|Zakres ubezpieczenia)",
        r"Agent\s+(.*?)\s+(?:Ubezpieczona działalność|Kod ryzyka|Zakres ubezpieczenia)",
        r"agent\s+(.*?)\s+ubezpieczona działalność",
    ], t)

    record.risk_code = find_first([
        r"Kod ryzyka(?:\s*/\s*risk code)?\s+([0-9]{5,6})",
        r"kod ryzyka\s+([0-9]{5,6})",
    ], t)

    record.insured_activity = extract_section(
        t,
        "Ubezpieczona działalność",
        [
            "Kod ryzyka",
            "Zakres ubezpieczenia",
            "Zakres ubezpieczenia/",
            "Suma gwarancyjna",
            "Klauzule dodatkowe",
        ],
    )

    record.scope_of_insurance = extract_section(
        t,
        "Zakres ubezpieczenia",
        [
            "Suma gwarancyjna",
            "Klauzule dodatkowe",
            "Franszyza redukcyjna",
            "prawo i jurysdykcja",
            "Okres ubezpieczenia",
            "Zakres terytorialny",
        ],
    )

    record.territorial_scope = find_first([
        r"Zakres terytorialny(?:\s*/\s*Territorial scope)?\s+(.*?)\s+(?:Okres ubezpieczenia|Składka|Franszyza redukcyjna|Płatność składki)",
        r"Zakres\s+Terytorialny\s+(.*?)\s+Okres",
        r"prawo i jurysdykcja\s+(.*?)\s+Okres ubezpieczenia",
    ], t)

    record.insurance_period = find_first([
        r"Okres ubezpieczenia(?:\s*/\s*Insurance period)?\s+(.*?)\s+(?:Składka|Stawka|Obrót|Płatność składki|Zasady i termin płatności)",
        r"Okres\s+Ubezpieczenia\s+(.*?)\s+Planowane",
    ], t)

    sum_block = extract_section(
        t,
        "Suma gwarancyjna",
        [
            "Franszyza redukcyjna",
            "Klauzule dodatkowe",
            "Klauzula dodatkowa",
            "Sublimity",
            "Zakres terytorialny",
            "Okres ubezpieczenia",
            "Planowany obrót",
            "Fundusz Płac",
            "Stawka",
            "Składka",
            "Płatność składki",
            "Prawo i jurysdykcja",
        ],
    )
    record.sum_guaranteed_notes = sum_block
    record.sum_guaranteed = compact_whitespace(sum_block)

    deductible_block = find_first([
        r"Franszyza redukcyjna(?:\s*/\s*Deductible)?\s+(.*?)\s+(?:prawo i jurysdykcja|Zakres terytorialny|Okres ubezpieczenia|Składka|Stawka|Płatność|Klauzula dodatkowa|Postanowienia dodatkowe)",
        r"franszyza redukcyjna\s+(.*?)\s+(?:zakres terytorialny|okres ubezpieczenia|składka|stawka|postanowienia dodatkowe)",
    ], t)
    record.deductible_notes = deductible_block
    record.deductible_main = find_first([
        r"([0-9\.\s]+(?:PLN|EUR|EURO|USD)\s+na każdą(?:\s+\w+){0,4})",
        r"([0-9\.\s]+(?:PLN|EUR|EURO|USD)\s+w każdej(?:\s+\w+){0,4})",
        r"([0-9\.\s]+(?:PLN|EUR|EURO|USD)\s+na każde roszczenie)",
    ], deductible_block)

    record.turnover = find_first([
        r"Obrót(?:\s*/\s*Turnover)?\s+(.*?)\s+(?:Stawka|Stopa składki|Składka)",
        r"Planowane\s+Obroty\s+(.*?)\s+Stawka",
    ], t)

    record.premium_rate = find_first([
        r"Stawka rozliczeniowa\s+(.*?)\s+(?:Składka|zasady i termin płatności|Zasady i termin płatności)",
        r"Stopa składki\s+(.*?)\s+Składka",
        r"Stawka\s+(.*?)\s+(?:Obrót|Składka|Planowane)",
    ], t)

    record.premium = find_first([
        r"Składka minimalna i zaliczkowa\s+(.*?)\s+(?:Zasady i termin płatności|Nr rachunku bankowego|Postanowienia dodatkowe)",
        r"Składka\s+(.*?)\s+(?:Płatność składki|Termin płatności składki|Nr rachunku bankowego)",
        r"składka\s+(.*?)\s+płatność składki",
    ], t)

    record.law_and_jurisdiction = find_first([
        r"prawo i jurysdykcja\s+(.*?)\s+(?:Klauzula dodatkowa|Okres ubezpieczenia|Stawka od obrotu|Stawka rozliczeniowa)",
        r"Prawo i jurysdykcja\s+(.*?)\s+Klauzula dodatkowa",
    ], t)

    record.has_office_liability = yesno("prowadzenia biura" in tl)
    record.has_documents_loss = yesno("zniszczenie i utrata dokumentów" in tl or "utrata dokumentów" in tl)
    record.has_subcontractors_cover = yesno("podwykonawc" in tl)

    regress_positive = ["zachowuje prawo regresu", "prawo regresu wobec podwykonawców"]
    regress_negative = ["rezygnuje z prawa regresu", "prawo regresu zostaje wyłączone"]
    if find_all_keywords(tl, regress_negative):
        record.has_subcontractors_regress = "limited_or_no"
    elif find_all_keywords(tl, regress_positive):
        record.has_subcontractors_regress = "yes"

    return record


def enrich_oc_ogolne(record: PolicyRecord, text: str) -> None:
    tl = lower_pl(text)
    record.covers_operations = yesno("prowadzonej działalności" in tl or "działalności gospodarczej" in tl)
    record.covers_product_liability = yesno("odpowiedzialność za produkt" in tl)
    record.covers_completed_operations = yesno("wykonane usługi" in tl or "completed operations" in tl)
    record.covers_professional_liability = "no"
    record.covers_pure_financial_loss = yesno("czyste szkody majątkowe" in tl)
    record.has_employer_liability = yesno("odpowiedzialność pracodawcy" in tl)
    record.has_property_under_control = yesno("rzeczy pod kontrolą" in tl or "mienie pod kontrolą" in tl)
    record.has_environment_damage = yesno("szkody w środowisku" in tl or "szkody ekologiczne" in tl)
    record.has_travel_clause = yesno("podróże służbowe" in tl)
    record.has_vehicles_clause = yesno("pojazdy nie podlegające obowiązkowi ubezpieczenia" in tl)
    record.has_leased_property = yesno("wziętych w najem" in tl or "najemcy" in tl)
    ext_prod_terms = [
        "klauzula maszynowa",
        "połączenia i zmieszania",
        "dalsze przetwarzanie i obróbkę",
        "kosztów usunięcia i zastąpienia",
        "wad etykiet",
    ]
    record.has_extended_product_clauses = yesno(find_all_keywords(tl, ext_prod_terms))


def enrich_oc_budowlane(record: PolicyRecord, text: str) -> None:
    tl = lower_pl(text)
    record.covers_professional_liability = "no"
    record.covers_pure_financial_loss = yesno("czyste szkody majątkowe" in tl)
