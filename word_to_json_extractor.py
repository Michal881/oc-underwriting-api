from __future__ import annotations

import json
import re
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document


@dataclass
class ExtractionResult:
    source_file: str
    source_type: str
    raw_text_excerpt: str
    fields: Dict[str, str]
    parsed_numbers: Dict[str, Optional[float]]
    flags: Dict[str, Optional[bool]]


FIELD_ALIASES: Dict[str, List[str]] = {
    "insurer": [
        "Ubezpieczyciel",
        "UBEZPIECZYCIEL/ INSURER",
        "Insurer",
    ],
    "policyholder": [
        "Ubezpieczający",
        "UBEZPIECZAJĄCY/ POLICYHOLDER",
        "Policyholder",
    ],
    "insured": [
        "Ubezpieczony",
        "Ubezpieczony/ insured",
        "insured",
    ],
    "agent_or_broker": [
        "Agent",
        "Broker",
    ],
    "insured_activity": [
        "Ubezpieczona działalność",
        "Ubezpieczona działalność/ insured business activity",
        "insured business activity",
    ],
    "insured_products": [
        "produkty ubezpieczone",
        "Produkty ubezpieczone",
    ],
    "risk_code": [
        "Kod ryzyka",
        "Kod ryzyka /risc code",
        "Kod ryzyka /risk code",
        "risk code",
        "risc code",
    ],
    "scope_of_insurance": [
        "Zakres ubezpieczenia",
        "Zakres ubezpieczenia /scope of insurance",
        "scope of insurance",
    ],
    "sum_guaranteed": [
        "Suma gwarancyjna",
        "Suma gwarancyjna/ sum insured",
        "sum insured",
    ],
    "additional_clauses": [
        "Klauzule dodatkowe",
        "Klauzule dodatkowe/",
        "additional clauses",
    ],
    "territorial_scope": [
        "Zakres terytorialny",
        "Zakres terytorialny /territorial scope",
        "territorial scope",
    ],
    "insurance_period": [
        "Okres ubezpieczenia",
        "Okres ubezpieczenia /insurance period",
        "insurance period",
    ],
    "turnover": [
        "Obrót",
        "obrót",
    ],
    "rate": [
        "Stawka",
        "stawka/ rate",
        "rate",
    ],
    "premium": [
        "Składka",
        "składka/premium",
        "premium",
        "składka",
    ],
    "premium_payment": [
        "Płatność składki",
        "Płatność składki/payment of the premium",
        "payment of the premium",
    ],
    "additional_conditions": [
        "Postanowienia dodatkowe",
        "postanowienia dodatkowe /additional conditions",
        "additional conditions",
    ],
    "reservation": [
        "Zastrzeżenie",
    ],
    "policy_conditions": [
        "Warunki OFERTY",
        "warunki polisy/ policy conditions",
        "policy conditions",
    ],
    "policy_number": [
        "Nr polisy",
        "Numer polisy",
        "Polisa nr",
        "POLISA Nr",
        "NR/No",
    ],
}

FIELD_ORDER = [
    "policy_number",
    "insurer",
    "policyholder",
    "insured",
    "agent_or_broker",
    "insured_activity",
    "insured_products",
    "risk_code",
    "scope_of_insurance",
    "sum_guaranteed",
    "additional_clauses",
    "territorial_scope",
    "insurance_period",
    "turnover",
    "rate",
    "premium",
    "premium_payment",
    "additional_conditions",
    "reservation",
    "policy_conditions",
]

MULTILINE_FIELDS = {
    "insured_activity",
    "insured_products",
    "scope_of_insurance",
    "additional_clauses",
    "premium_payment",
    "additional_conditions",
    "reservation",
    "policy_conditions",
}

SIMPLE_FIELDS = {
    "insurer",
    "policyholder",
    "insured",
    "agent_or_broker",
    "risk_code",
    "sum_guaranteed",
    "territorial_scope",
    "insurance_period",
    "turnover",
    "rate",
    "premium",
    "policy_number",
}


def clean_text(text: str) -> str:
    text = text.replace("\xa0", " ")
    text = text.replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def normalize_for_matching(text: str) -> str:
    text = clean_text(text).lower()
    text = text.replace(":", " ")
    text = text.replace("/", " ")
    text = text.replace("|", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def dedupe_preserve_order(items: List[str]) -> List[str]:
    seen = set()
    out = []
    for item in items:
        norm = normalize_for_matching(item)
        if norm and norm not in seen:
            seen.add(norm)
            out.append(item.strip())
    return out


def read_docx_text(file_path: Path) -> str:
    document = Document(str(file_path))
    chunks: List[str] = []

    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)
        if text:
            chunks.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                chunks.append(" | ".join(cells))

    return "\n".join(chunks)


def build_alias_index() -> Dict[str, str]:
    alias_to_field: Dict[str, str] = {}
    for field_name, aliases in FIELD_ALIASES.items():
        for alias in aliases:
            alias_to_field[normalize_for_matching(alias)] = field_name
    return alias_to_field


def split_table_line(line: str) -> List[str]:
    parts = [clean_text(part) for part in line.split("|")]
    parts = [part for part in parts if part]
    return dedupe_preserve_order(parts)


def find_field_in_text(line: str, alias_index: Dict[str, str]) -> Optional[str]:
    normalized_line = normalize_for_matching(line)
    for alias_norm, field_name in alias_index.items():
        if normalized_line == alias_norm:
            return field_name
    return None


def find_field_in_table_parts(parts: List[str], alias_index: Dict[str, str]) -> Optional[str]:
    for part in parts:
        normalized = normalize_for_matching(part)
        if normalized in alias_index:
            return alias_index[normalized]
    return None


def contains_known_field_label(text: str, alias_index: Dict[str, str]) -> bool:
    normalized_text = normalize_for_matching(text)
    return normalized_text in alias_index


def looks_like_policy_number(text: str) -> bool:
    text = clean_text(text)
    return bool(re.search(r"\b\d+/\d+/\d+/\d+/\d+\b", text))


def extract_policy_number_from_text(raw_text: str) -> Optional[str]:
    patterns = [
        r"NR/No\s+(\d+/\d+/\d+/\d+/\d+)",
        r"POLISA\s+Nr\s+(\d+/\d+/\d+/\d+/\d+)",
        r"Polisa\s+nr\s+(\d+/\d+/\d+/\d+/\d+)",
    ]
    for pattern in patterns:
        match = re.search(pattern, raw_text, flags=re.IGNORECASE)
        if match:
            return match.group(1)
    return None


def split_polish_from_english(text: str) -> str:
    text = clean_text(text)

    markers = [
        " / Production of",
        " / General",
        " / Worldwide",
        " / each and every",
        " / each loss",
        " / each Event",
        " / Extended",
        " / Deductible",
        " / Liability",
        " / Insurance",
        " / Policyholder",
        " / insured",
        " / insurer",
        " / sum insured",
        " /scope of insurance",
        " /territorial scope",
        " /insurance period",
        " /risk code",
        " /risc code",
    ]

    min_idx = None
    for marker in markers:
        idx = text.find(marker)
        if idx != -1 and (min_idx is None or idx < min_idx):
            min_idx = idx

    if min_idx is not None:
        text = text[:min_idx]

    return text.strip(" /-")


def normalize_multiline_value(text: str) -> str:
    lines = [clean_text(x) for x in text.splitlines()]
    lines = [x for x in lines if x]
    lines = dedupe_preserve_order(lines)
    return "\n".join(lines)


def extract_money_strings(text: str) -> List[str]:
    return re.findall(
        r"\d{1,3}(?:[ .]\d{3})*(?:,\d+)?\s*(?:PLN|EUR|USD)",
        text,
        flags=re.IGNORECASE,
    )


def extract_rate_strings(text: str) -> List[str]:
    return re.findall(r"\d+(?:,\d+)?\s*‰", text)


def postprocess_field_value(field_name: str, value: str, raw_text: str = "") -> str:
    value = clean_text(value)

    if field_name == "policy_number":
        if looks_like_policy_number(value):
            match = re.search(r"\b\d+/\d+/\d+/\d+/\d+\b", value)
            if match:
                return match.group(0)
        from_text = extract_policy_number_from_text(raw_text)
        return from_text or ""

    if field_name == "risk_code":
        match = re.search(r"\b\d{5,8}\b", value)
        if match:
            return match.group(0)
        return ""

    if field_name == "sum_guaranteed":
        match = re.search(
            r"\d{1,3}(?:[ .]\d{3})*(?:,\d+)?\s*(?:PLN|EUR|USD).*",
            value,
            flags=re.IGNORECASE,
        )
        if match:
            return clean_text(match.group(0))

    if field_name == "turnover":
        match = re.search(
            r"\d{1,3}(?:[ .]\d{3})*(?:,\d+)?\s*(?:PLN|EUR|USD).*",
            value,
            flags=re.IGNORECASE,
        )
        if match:
            return clean_text(match.group(0))

    if field_name == "premium":
        amounts = extract_money_strings(value)
        if amounts:
            return " | ".join(dedupe_preserve_order(amounts))

    if field_name == "rate":
        rates = extract_rate_strings(value)
        if rates:
            return " | ".join(dedupe_preserve_order(rates))

    if field_name == "insurance_period":
        match = re.search(
            r"Od dnia\s+\d{2}[.\-]\d{2}[.\-]\d{4}\s*r?\.\s*do dnia\s+\d{2}[.\-]\d{2}[.\-]\d{4}\s*r?\.",
            value,
            flags=re.IGNORECASE,
        )
        if match:
            return clean_text(match.group(0))

        match = re.search(
            r"\d{2}[.\-]\d{2}[.\-]\d{4}\s*[–-]\s*\d{2}[.\-]\d{2}[.\-]\d{4}",
            value
        )
        if match:
            return match.group(0)

    if field_name in {
        "insurer",
        "policyholder",
        "insured",
        "agent_or_broker",
        "insured_activity",
        "insured_products",
        "territorial_scope",
    }:
        value = split_polish_from_english(value)

    value = re.sub(
        r"^(pieczyciel|pieczający|pieczony|eczona działalność|es ubezpieczenia|warancyjna|terytorialny|ienia dodatkowe)\s*\|\s*",
        "",
        value,
        flags=re.IGNORECASE,
    )
    value = re.sub(
        r"\b(Ubezpieczyciel|Ubezpieczający|Ubezpieczony|Broker|Ubezpieczona działalność|Zakres ubezpieczenia|Suma gwarancyjna|Klauzule dodatkowe|Zakres terytorialny|Płatność składki|Postanowienia dodatkowe|Składka)\b\s*\|\s*",
        "",
        value,
        flags=re.IGNORECASE,
    )

    parts = split_table_line(value)
    parts = [
        p for p in parts
        if not re.fullmatch(
            r"(Ubezpieczyciel|Ubezpieczający|Ubezpieczony|Broker|Ubezpieczona działalność|Zakres ubezpieczenia|Suma gwarancyjna|Klauzule dodatkowe|Zakres terytorialny|Płatność składki|Postanowienia dodatkowe|Składka)",
            p,
            flags=re.IGNORECASE,
        )
    ]
    parts = dedupe_preserve_order(parts)

    if field_name in MULTILINE_FIELDS:
        value = "\n".join(parts) if parts else value
        return normalize_multiline_value(value)

    if len(parts) == 1:
        return parts[0]

    if len(parts) > 1:
        return " | ".join(parts)

    return value


def is_continuation_line_for_field(
    line: str,
    current_field: str,
    alias_index: Dict[str, str],
) -> bool:
    line = clean_text(line)
    if not line:
        return False

    parts = split_table_line(line)

    if find_field_in_table_parts(parts, alias_index):
        return False

    if find_field_in_text(line, alias_index):
        return False

    if current_field in MULTILINE_FIELDS:
        return True

    if current_field == "insurance_period":
        return "do dnia" in line.lower() or bool(re.search(r"\d{2}[.\-]\d{2}[.\-]\d{4}", line))

    if current_field == "premium":
        if "płatna" in line.lower():
            return False
        return bool(extract_money_strings(line))

    if current_field in {"turnover", "sum_guaranteed", "rate"}:
        return bool(extract_money_strings(line) or extract_rate_strings(line))

    return False


def append_field_value(fields: Dict[str, str], field_name: str, value: str) -> None:
    value = clean_text(value)
    if not value:
        return

    if field_name in fields and fields[field_name]:
        existing_lines = [clean_text(x) for x in fields[field_name].splitlines() if clean_text(x)]
        new_lines = [clean_text(x) for x in value.splitlines() if clean_text(x)]
        merged = dedupe_preserve_order(existing_lines + new_lines)
        fields[field_name] = "\n".join(merged)
    else:
        fields[field_name] = value


def extract_fields_from_lines(lines: List[str], raw_text: str) -> Dict[str, str]:
    alias_index = build_alias_index()
    fields: Dict[str, str] = {}

    current_field: Optional[str] = None
    current_buffer: List[str] = []

    def flush() -> None:
        nonlocal current_field, current_buffer
        if current_field and current_buffer:
            value = clean_text("\n".join(current_buffer))
            value = postprocess_field_value(current_field, value, raw_text=raw_text)
            append_field_value(fields, current_field, value)
        current_field = None
        current_buffer = []

    i = 0
    while i < len(lines):
        line = clean_text(lines[i])
        if not line:
            i += 1
            continue

        parts = split_table_line(line)
        table_field = find_field_in_table_parts(parts, alias_index)

        if table_field:
            flush()

            raw_value_parts = [
                p for p in parts
                if not contains_known_field_label(p, alias_index)
            ]
            raw_value_parts = dedupe_preserve_order(raw_value_parts)
            initial_value = "\n".join(raw_value_parts).strip()

            current_field = table_field
            current_buffer = [initial_value] if initial_value else []

            j = i + 1
            while j < len(lines):
                next_line = clean_text(lines[j])
                if not next_line:
                    j += 1
                    continue

                if is_continuation_line_for_field(next_line, current_field, alias_index):
                    current_buffer.append(next_line)
                    j += 1
                    continue

                break

            flush()
            i = j
            continue

        plain_field = find_field_in_text(line, alias_index)
        if plain_field:
            flush()
            current_field = plain_field
            current_buffer = []
            i += 1
            continue

        if current_field and is_continuation_line_for_field(line, current_field, alias_index):
            current_buffer.append(line)
            i += 1
            continue

        flush()
        i += 1

    flush()

    if "policy_number" not in fields:
        policy_number = extract_policy_number_from_text(raw_text)
        if policy_number:
            fields["policy_number"] = policy_number

    # Dodatkowy fallback dla premium
    if "premium" not in fields:
        premium_candidates = re.findall(
            r"Składka.*?(\d{1,3}(?:[ .]\d{3})*(?:,\d+)?\s*(?:PLN|EUR|USD))",
            raw_text,
            flags=re.IGNORECASE | re.DOTALL,
        )
        if premium_candidates:
            fields["premium"] = " | ".join(dedupe_preserve_order([clean_text(x) for x in premium_candidates]))

    return fields


def parse_money_like(value: Optional[str]) -> Optional[float]:
    if not value:
        return None
    match = re.search(r"(\d{1,3}(?:[ .]\d{3})*(?:,\d+)?)", value)
    if not match:
        return None
    number = match.group(1).replace(" ", "").replace(".", "").replace(",", ".")
    try:
        return float(number)
    except ValueError:
        return None


def parse_per_mille(value: Optional[str]) -> Optional[float]:
    if not value:
        return None
    match = re.search(r"(\d+(?:,\d+)?)\s*‰", value)
    if not match:
        return None
    return float(match.group(1).replace(",", "."))


def contains_any(text: str, needles: List[str]) -> bool:
    text = text.lower()
    return any(needle.lower() in text for needle in needles)


def build_flags(fields: Dict[str, str]) -> Dict[str, Optional[bool]]:
    scope = "\n".join([
        fields.get("scope_of_insurance", ""),
        fields.get("additional_clauses", ""),
        fields.get("additional_conditions", ""),
    ]).lower()

    territorial = fields.get("territorial_scope", "").lower()

    return {
        "covers_professional_liability": contains_any(scope, ["zawodowa", "wykonywania zawodu"]),
        "covers_general_liability": contains_any(scope, ["prowadzonej działalności", "posiadanego mienia", "third party liability"]),
        "covers_product_liability": contains_any(scope, ["produkt", "product liability"]),
        "covers_completed_operations": contains_any(scope, ["wykonaną usługę", "completed operations"]),
        "covers_pure_financial_loss": contains_any(scope, ["czyste szkody majątkowe", "pure financial losses"]),
        "covers_employers_liability": contains_any(scope, ["odpowiedzialność cywilna pracodawcy", "odpowiedzialność pracodawcy", "employer’s liability", "employers liability"]),
        "covers_documents_loss": contains_any(scope, ["utrata dokumentów", "zniszczenie i utrata dokumentów"]),
        "covers_subcontractors": contains_any(scope, ["podwykonawców", "subcontractors"]),
        "usa_canada_included": contains_any(territorial + "\n" + scope, ["usa i kanady", "usa and canada", "kanady"]),
    }


def parse_numbers(fields: Dict[str, str]) -> Dict[str, Optional[float]]:
    return {
        "turnover_amount": parse_money_like(fields.get("turnover")),
        "sum_guaranteed_amount": parse_money_like(fields.get("sum_guaranteed")),
        "premium_amount": parse_money_like(fields.get("premium")),
        "rate_primary_per_mille": parse_per_mille(fields.get("rate")),
    }


def extract_policy(file_path: Path) -> ExtractionResult:
    raw_text = read_docx_text(file_path)
    lines = raw_text.splitlines()
    fields = extract_fields_from_lines(lines, raw_text)
    parsed_numbers = parse_numbers(fields)
    flags = build_flags(fields)

    ordered_fields = {k: fields[k] for k in FIELD_ORDER if k in fields}

    return ExtractionResult(
        source_file=file_path.name,
        source_type=file_path.suffix.lower().lstrip("."),
        raw_text_excerpt=raw_text[:5000],
        fields=ordered_fields,
        parsed_numbers=parsed_numbers,
        flags=flags,
    )


def save_json(result: ExtractionResult, output_path: Path) -> None:
    output_path.write_text(
        json.dumps(asdict(result), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def main() -> None:
    import argparse

    parser = argparse.ArgumentParser(description="Ekstrakcja danych z polis DOCX do JSON (batch)")
    parser.add_argument(
        "-d",
        "--directory",
        help="Folder z plikami .docx",
        required=True,
    )
    args = parser.parse_args()

    input_dir = Path(args.directory)

    if not input_dir.exists():
        raise FileNotFoundError(f"Nie znaleziono folderu: {input_dir}")

    files = sorted(list(input_dir.glob("*.docx")))

    print(f"Znaleziono {len(files)} plików...")

    for file_path in files:
        try:
            result = extract_policy(file_path)
            output_path = file_path.with_suffix(".json")
            save_json(result, output_path)

            print(f"\n✔ {file_path.name} → OK")
            print(f"  risk_code: {result.fields.get('risk_code')}")
            print(f"  sum_guaranteed: {result.fields.get('sum_guaranteed')}")
            print(f"  turnover: {result.fields.get('turnover')}")
            print(f"  rate: {result.fields.get('rate')}")
            print(f"  premium: {result.fields.get('premium')}")

        except Exception as e:
            print(f"\n❌ {file_path.name} → BŁĄD: {e}")


if __name__ == "__main__":
    main()